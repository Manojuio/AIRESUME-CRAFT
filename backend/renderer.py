import os
import logging
from pathlib import Path

from jinja2 import Environment, FileSystemLoader
from docx import Document
from docx.shared import Pt

logger = logging.getLogger(__name__)

GENERATED_DIR = Path(__file__).parent / "generated"
TEMPLATE_DIR = Path(__file__).parent / "templates"
env = Environment(loader=FileSystemLoader(str(TEMPLATE_DIR)))


def render_pdf(resume_data: dict, task_id: str) -> str:
    try:
        from weasyprint import HTML

        template = env.get_template("resume.html")
        html = template.render(resume=resume_data)

        task_dir = GENERATED_DIR / task_id
        task_dir.mkdir(parents=True, exist_ok=True)
        pdf_path = task_dir / "resume.pdf"
        HTML(string=html).write_pdf(str(pdf_path))
        return str(pdf_path)
    except Exception as e:
        logger.warning(f"PDF rendering failed, skipping PDF generation: {e}")
        # Return None or empty path if PDF fails - frontend will handle it
        return None


def render_docx(resume_data: dict, task_id: str) -> str:
    try:
        doc = Document()
        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(11)

        doc.add_heading(resume_data.get("full_name", "Resume"), level=1)

        email = resume_data.get("email", "")
        if email:
            doc.add_paragraph(f"Email: {email}")

        summary = resume_data.get("summary", "")
        if summary:
            doc.add_heading("Summary", level=2)
            doc.add_paragraph(summary)

        skills = resume_data.get("skills", [])
        if skills:
            doc.add_heading("Skills", level=2)
            for skill in skills:
                doc.add_paragraph(skill, style="List Bullet")

        experience = resume_data.get("experience", [])
        if experience:
            doc.add_heading("Experience", level=2)
            for exp in experience:
                p = doc.add_paragraph()
                run = p.add_run(exp.get("project_name", ""))
                run.bold = True
                if exp.get("technologies"):
                    p.add_run(f" — {exp['technologies']}")
                doc.add_paragraph(exp.get("description", ""))

        education = resume_data.get("education", [])
        if education:
            doc.add_heading("Education", level=2)
            for edu in education:
                p = doc.add_paragraph()
                run = p.add_run(edu.get("degree", ""))
                run.bold = True
                p.add_run(f" — {edu.get('institution', '')} ({edu.get('year', '')})")

        task_dir = GENERATED_DIR / task_id
        task_dir.mkdir(parents=True, exist_ok=True)
        docx_path = task_dir / "resume.docx"
        doc.save(str(docx_path))
        return str(docx_path)
    except Exception as e:
        logger.warning(f"DOCX rendering failed: {e}")
        return None


def render_portfolio_html(resume_data: dict, top5_repos: list, task_id: str) -> str:
    try:
        template = env.get_template("portfolio.html")
        html = template.render(resume=resume_data, repos=top5_repos)

        task_dir = GENERATED_DIR / task_id
        task_dir.mkdir(parents=True, exist_ok=True)
        portfolio_path = task_dir / "portfolio.html"
        portfolio_path.write_text(html, encoding="utf-8")
        return str(portfolio_path)
    except Exception as e:
        logger.warning(f"Portfolio HTML rendering failed: {e}")
        return None
